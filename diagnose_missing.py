from docx import Document

TARGET_TITLES = [
    "the man who desired gold",
    "the richest man in babylon",
]

DOC_PATH = "The Richest Man In Babylon_Full_Book_Source.docx"

doc = Document(DOC_PATH)

print("=" * 70)
print("SEARCHING FOR MISSING CHAPTER HEADINGS")
print("=" * 70)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    text_lower = text.lower()
    for target in TARGET_TITLES:
        if target in text_lower:
            runs_info = []
            for r in para.runs:
                if r.text.strip():
                    runs_info.append({
                        "text": repr(r.text),
                        "bold": r.bold,
                        "italic": r.italic,
                        "font_size": r.font.size,
                        "font_name": r.font.name,
                    })

            print(f"\npara[{i:04d}]  TEXT: {repr(text)}")
            print(f"           style: {para.style.name!r}")
            print(f"           all_caps check: {text == text.upper()}")
            print(f"           runs ({len(para.runs)} total):")
            for r in runs_info:
                print(f"             {r}")
            break

print("\n" + "=" * 70)
print("ALSO SHOWING paras 0..60 with bold/italic runs")
print("=" * 70)

for i, para in enumerate(doc.paragraphs[:60]):
    text = para.text.strip()
    if not text:
        continue
    has_bold   = any(r.bold   for r in para.runs if r.text.strip())
    has_italic = any(r.italic for r in para.runs if r.text.strip())
    if has_bold or has_italic:
        print(f"para[{i:04d}]  bold={has_bold}  italic={has_italic}  style={para.style.name!r}  TEXT={repr(text[:80])}")